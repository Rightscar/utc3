"""
Multi-Format Export Engine - Phase 4
====================================

Comprehensive export system supporting multiple formats:
- JSON/JSONL for APIs and data processing
- CSV for spreadsheet analysis
- Markdown for documentation
- TXT for simple text files
- DOCX for Word documents
- PDF for professional reports
- ZIP archives for bulk exports
"""

import streamlit as st
import logging
import json
import csv
import io
import zipfile
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, asdict
from datetime import datetime
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ExportRequest:
    """Represents an export request"""
    content_items: List[Any]
    export_format: str
    filename: str
    metadata: Dict[str, Any]
    include_metadata: bool = True
    compression: bool = False

@dataclass
class ExportResult:
    """Represents the result of an export operation"""
    success: bool
    filename: str
    file_size: int
    export_format: str
    items_exported: int
    download_data: bytes
    error_message: str = ""

class MultiFormatExporter:
    """Engine for exporting content in multiple formats"""
    
    def __init__(self):
        self.supported_formats = self._load_supported_formats()
        
    def _load_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Load supported export formats and their configurations"""
        return {
            "json": {
                "name": "JSON",
                "description": "JavaScript Object Notation - structured data format",
                "mime_type": "application/json",
                "file_extension": ".json",
                "supports_metadata": True,
                "supports_compression": True,
                "use_cases": ["API integration", "Data processing", "Web applications"]
            },
            
            "jsonl": {
                "name": "JSONL",
                "description": "JSON Lines - one JSON object per line",
                "mime_type": "application/jsonl",
                "file_extension": ".jsonl",
                "supports_metadata": True,
                "supports_compression": True,
                "use_cases": ["Machine learning", "Streaming data", "Large datasets"]
            },
            
            "csv": {
                "name": "CSV",
                "description": "Comma-Separated Values - spreadsheet format",
                "mime_type": "text/csv",
                "file_extension": ".csv",
                "supports_metadata": False,
                "supports_compression": True,
                "use_cases": ["Excel analysis", "Data visualization", "Statistical analysis"]
            },
            
            "markdown": {
                "name": "Markdown",
                "description": "Markdown format - human-readable documentation",
                "mime_type": "text/markdown",
                "file_extension": ".md",
                "supports_metadata": True,
                "supports_compression": True,
                "use_cases": ["Documentation", "GitHub README", "Blog posts"]
            },
            
            "txt": {
                "name": "Plain Text",
                "description": "Simple text format - universal compatibility",
                "mime_type": "text/plain",
                "file_extension": ".txt",
                "supports_metadata": False,
                "supports_compression": True,
                "use_cases": ["Simple sharing", "Email content", "Basic storage"]
            },
            
            "docx": {
                "name": "Word Document",
                "description": "Microsoft Word format - professional documents",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "file_extension": ".docx",
                "supports_metadata": True,
                "supports_compression": False,
                "use_cases": ["Professional reports", "Formatted documents", "Collaboration"]
            },
            
            "pdf": {
                "name": "PDF",
                "description": "Portable Document Format - professional presentation",
                "mime_type": "application/pdf",
                "file_extension": ".pdf",
                "supports_metadata": True,
                "supports_compression": False,
                "use_cases": ["Professional sharing", "Print-ready documents", "Archival"]
            },
            
            "zip": {
                "name": "ZIP Archive",
                "description": "Compressed archive with multiple formats",
                "mime_type": "application/zip",
                "file_extension": ".zip",
                "supports_metadata": True,
                "supports_compression": True,
                "use_cases": ["Bulk export", "Multiple formats", "Backup"]
            }
        }
    
    def get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Get all supported export formats"""
        return self.supported_formats
    
    def get_format_recommendations(self, use_case: str) -> List[str]:
        """Get format recommendations based on use case"""
        
        recommendations = {
            "api_integration": ["json", "jsonl"],
            "data_analysis": ["csv", "json"],
            "documentation": ["markdown", "pdf", "docx"],
            "sharing": ["pdf", "docx", "txt"],
            "machine_learning": ["jsonl", "json", "csv"],
            "backup": ["zip", "json"],
            "web_development": ["json", "markdown"],
            "research": ["pdf", "docx", "csv"],
            "general": ["json", "markdown", "txt"]
        }
        
        return recommendations.get(use_case, recommendations["general"])
    
    def export_content(self, content_items: List[Any], export_format: str,
                      filename: str = None, **kwargs) -> ExportResult:
        """
        Export content in specified format
        
        Args:
            content_items: List of content items to export
            export_format: Target export format
            filename: Optional custom filename
            **kwargs: Additional export options
            
        Returns:
            ExportResult object
        """
        if export_format not in self.supported_formats:
            return ExportResult(
                success=False,
                filename="",
                file_size=0,
                export_format=export_format,
                items_exported=0,
                download_data=b"",
                error_message=f"Unsupported format: {export_format}"
            )
        
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                extension = self.supported_formats[export_format]["file_extension"]
                filename = f"exported_content_{timestamp}{extension}"
            
            # Export based on format
            if export_format == "json":
                download_data = self._export_json(content_items, **kwargs)
            elif export_format == "jsonl":
                download_data = self._export_jsonl(content_items, **kwargs)
            elif export_format == "csv":
                download_data = self._export_csv(content_items, **kwargs)
            elif export_format == "markdown":
                download_data = self._export_markdown(content_items, **kwargs)
            elif export_format == "txt":
                download_data = self._export_txt(content_items, **kwargs)
            elif export_format == "docx":
                download_data = self._export_docx(content_items, **kwargs)
            elif export_format == "pdf":
                download_data = self._export_pdf(content_items, **kwargs)
            elif export_format == "zip":
                download_data = self._export_zip(content_items, **kwargs)
            else:
                raise ValueError(f"Export method not implemented for {export_format}")
            
            return ExportResult(
                success=True,
                filename=filename,
                file_size=len(download_data),
                export_format=export_format,
                items_exported=len(content_items),
                download_data=download_data
            )
            
        except Exception as e:
            logger.error(f"Export error: {e}")
            return ExportResult(
                success=False,
                filename=filename or "",
                file_size=0,
                export_format=export_format,
                items_exported=0,
                download_data=b"",
                error_message=str(e)
            )
    
    def _export_json(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as JSON"""
        
        include_metadata = kwargs.get('include_metadata', True)
        
        export_data = {
            "metadata": {
                "export_timestamp": datetime.now().isoformat(),
                "total_items": len(content_items),
                "export_format": "json",
                "generator": "Universal Content Generator MVP"
            } if include_metadata else {},
            "content": []
        }
        
        for item in content_items:
            if hasattr(item, '__dict__'):
                # Convert dataclass or object to dict
                item_dict = asdict(item) if hasattr(item, '__dataclass_fields__') else vars(item)
            elif isinstance(item, dict):
                item_dict = item
            else:
                item_dict = {"content": str(item)}
            
            export_data["content"].append(item_dict)
        
        json_str = json.dumps(export_data, indent=2, ensure_ascii=False, default=str)
        return json_str.encode('utf-8')
    
    def _export_jsonl(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as JSONL (JSON Lines)"""
        
        lines = []
        
        for item in content_items:
            if hasattr(item, '__dict__'):
                item_dict = asdict(item) if hasattr(item, '__dataclass_fields__') else vars(item)
            elif isinstance(item, dict):
                item_dict = item
            else:
                item_dict = {"content": str(item)}
            
            line = json.dumps(item_dict, ensure_ascii=False, default=str)
            lines.append(line)
        
        jsonl_str = '\n'.join(lines)
        return jsonl_str.encode('utf-8')
    
    def _export_csv(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as CSV"""
        
        if not content_items:
            return b""
        
        output = io.StringIO()
        
        # Determine CSV structure based on first item
        first_item = content_items[0]
        
        if hasattr(first_item, '__dict__'):
            # Object with attributes
            fieldnames = list(vars(first_item).keys())
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            
            for item in content_items:
                row_data = vars(item)
                # Convert complex objects to strings
                for key, value in row_data.items():
                    if isinstance(value, (list, dict)):
                        row_data[key] = json.dumps(value, default=str)
                    elif not isinstance(value, (str, int, float, bool, type(None))):
                        row_data[key] = str(value)
                writer.writerow(row_data)
        
        elif isinstance(first_item, dict):
            # Dictionary items
            fieldnames = list(first_item.keys())
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            
            for item in content_items:
                row_data = item.copy()
                for key, value in row_data.items():
                    if isinstance(value, (list, dict)):
                        row_data[key] = json.dumps(value, default=str)
                    elif not isinstance(value, (str, int, float, bool, type(None))):
                        row_data[key] = str(value)
                writer.writerow(row_data)
        
        else:
            # Simple items
            writer = csv.writer(output)
            writer.writerow(["content"])
            for item in content_items:
                writer.writerow([str(item)])
        
        csv_str = output.getvalue()
        output.close()
        
        return csv_str.encode('utf-8')
    
    def _export_markdown(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as Markdown"""
        
        include_metadata = kwargs.get('include_metadata', True)
        title = kwargs.get('title', 'Exported Content')
        
        markdown_lines = []
        
        # Header
        markdown_lines.append(f"# {title}")
        markdown_lines.append("")
        
        # Metadata
        if include_metadata:
            markdown_lines.append("## Export Information")
            markdown_lines.append("")
            markdown_lines.append(f"- **Export Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            markdown_lines.append(f"- **Total Items:** {len(content_items)}")
            markdown_lines.append(f"- **Format:** Markdown")
            markdown_lines.append("")
            markdown_lines.append("---")
            markdown_lines.append("")
        
        # Content
        markdown_lines.append("## Content")
        markdown_lines.append("")
        
        for i, item in enumerate(content_items, 1):
            markdown_lines.append(f"### Item {i}")
            markdown_lines.append("")
            
            if hasattr(item, 'title') and item.title:
                markdown_lines.append(f"**Title:** {item.title}")
                markdown_lines.append("")
            
            # Main content
            if hasattr(item, 'rewritten_text'):
                markdown_lines.append(item.rewritten_text)
            elif hasattr(item, 'narrated_content'):
                markdown_lines.append(item.narrated_content)
            elif hasattr(item, 'content'):
                markdown_lines.append(item.content)
            elif hasattr(item, 'input_text') and hasattr(item, 'output_text'):
                markdown_lines.append(f"**Input:** {item.input_text}")
                markdown_lines.append("")
                markdown_lines.append(f"**Output:** {item.output_text}")
            else:
                markdown_lines.append(str(item))
            
            markdown_lines.append("")
            
            # Metadata for item
            if hasattr(item, 'quality_score'):
                markdown_lines.append(f"*Quality Score: {item.quality_score:.2f}*")
                markdown_lines.append("")
            
            markdown_lines.append("---")
            markdown_lines.append("")
        
        markdown_str = '\n'.join(markdown_lines)
        return markdown_str.encode('utf-8')
    
    def _export_txt(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as plain text"""
        
        lines = []
        
        lines.append("EXPORTED CONTENT")
        lines.append("=" * 50)
        lines.append("")
        lines.append(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Total Items: {len(content_items)}")
        lines.append("")
        lines.append("=" * 50)
        lines.append("")
        
        for i, item in enumerate(content_items, 1):
            lines.append(f"ITEM {i}")
            lines.append("-" * 20)
            lines.append("")
            
            # Extract main content
            if hasattr(item, 'rewritten_text'):
                lines.append(item.rewritten_text)
            elif hasattr(item, 'narrated_content'):
                lines.append(item.narrated_content)
            elif hasattr(item, 'content'):
                lines.append(item.content)
            elif hasattr(item, 'input_text') and hasattr(item, 'output_text'):
                lines.append(f"INPUT: {item.input_text}")
                lines.append("")
                lines.append(f"OUTPUT: {item.output_text}")
            else:
                lines.append(str(item))
            
            lines.append("")
            lines.append("")
        
        txt_str = '\n'.join(lines)
        return txt_str.encode('utf-8')
    
    def _export_docx(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as Word document"""
        
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            # Fallback to text export if python-docx not available
            logger.warning("python-docx not available, falling back to text export")
            return self._export_txt(content_items, **kwargs)
        
        doc = Document()
        
        # Title
        title = kwargs.get('title', 'Exported Content')
        doc.add_heading(title, 0)
        
        # Metadata
        if kwargs.get('include_metadata', True):
            doc.add_heading('Export Information', level=1)
            p = doc.add_paragraph()
            p.add_run(f"Export Date: ").bold = True
            p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            p = doc.add_paragraph()
            p.add_run(f"Total Items: ").bold = True
            p.add_run(str(len(content_items)))
            
            doc.add_page_break()
        
        # Content
        doc.add_heading('Content', level=1)
        
        for i, item in enumerate(content_items, 1):
            doc.add_heading(f'Item {i}', level=2)
            
            # Main content
            if hasattr(item, 'rewritten_text'):
                doc.add_paragraph(item.rewritten_text)
            elif hasattr(item, 'narrated_content'):
                doc.add_paragraph(item.narrated_content)
            elif hasattr(item, 'content'):
                doc.add_paragraph(item.content)
            elif hasattr(item, 'input_text') and hasattr(item, 'output_text'):
                p = doc.add_paragraph()
                p.add_run("Input: ").bold = True
                p.add_run(item.input_text)
                
                p = doc.add_paragraph()
                p.add_run("Output: ").bold = True
                p.add_run(item.output_text)
            else:
                doc.add_paragraph(str(item))
            
            # Quality score if available
            if hasattr(item, 'quality_score'):
                p = doc.add_paragraph()
                p.add_run(f"Quality Score: {item.quality_score:.2f}").italic = True
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    def _export_pdf(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as PDF"""
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        except ImportError:
            # Fallback to text export if reportlab not available
            logger.warning("reportlab not available, falling back to text export")
            return self._export_txt(content_items, **kwargs)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
        )
        
        story = []
        
        # Title
        title = kwargs.get('title', 'Exported Content')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Metadata
        if kwargs.get('include_metadata', True):
            story.append(Paragraph('Export Information', heading_style))
            story.append(Paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Paragraph(f"Total Items: {len(content_items)}", styles['Normal']))
            story.append(PageBreak())
        
        # Content
        story.append(Paragraph('Content', heading_style))
        story.append(Spacer(1, 12))
        
        for i, item in enumerate(content_items, 1):
            story.append(Paragraph(f'Item {i}', heading_style))
            
            # Main content
            content_text = ""
            if hasattr(item, 'rewritten_text'):
                content_text = item.rewritten_text
            elif hasattr(item, 'narrated_content'):
                content_text = item.narrated_content
            elif hasattr(item, 'content'):
                content_text = item.content
            elif hasattr(item, 'input_text') and hasattr(item, 'output_text'):
                content_text = f"<b>Input:</b> {item.input_text}<br/><br/><b>Output:</b> {item.output_text}"
            else:
                content_text = str(item)
            
            story.append(Paragraph(content_text, styles['Normal']))
            
            # Quality score if available
            if hasattr(item, 'quality_score'):
                story.append(Paragraph(f"<i>Quality Score: {item.quality_score:.2f}</i>", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _export_zip(self, content_items: List[Any], **kwargs) -> bytes:
        """Export content as ZIP archive with multiple formats"""
        
        buffer = io.BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Export in multiple formats
            formats_to_include = kwargs.get('formats', ['json', 'markdown', 'txt'])
            
            for format_name in formats_to_include:
                if format_name in self.supported_formats and format_name != 'zip':
                    try:
                        # Generate content for this format
                        if format_name == 'json':
                            content_data = self._export_json(content_items, **kwargs)
                        elif format_name == 'markdown':
                            content_data = self._export_markdown(content_items, **kwargs)
                        elif format_name == 'txt':
                            content_data = self._export_txt(content_items, **kwargs)
                        elif format_name == 'csv':
                            content_data = self._export_csv(content_items, **kwargs)
                        elif format_name == 'jsonl':
                            content_data = self._export_jsonl(content_items, **kwargs)
                        else:
                            continue
                        
                        # Add to ZIP
                        extension = self.supported_formats[format_name]['file_extension']
                        filename = f"exported_content{extension}"
                        zip_file.writestr(filename, content_data)
                        
                    except Exception as e:
                        logger.error(f"Error adding {format_name} to ZIP: {e}")
            
            # Add metadata file
            metadata = {
                "export_timestamp": datetime.now().isoformat(),
                "total_items": len(content_items),
                "formats_included": formats_to_include,
                "generator": "Universal Content Generator MVP"
            }
            
            metadata_json = json.dumps(metadata, indent=2)
            zip_file.writestr("export_metadata.json", metadata_json)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def create_streamlit_download(self, export_result: ExportResult) -> None:
        """Create Streamlit download button for export result"""
        
        if not export_result.success:
            st.error(f"Export failed: {export_result.error_message}")
            return
        
        format_info = self.supported_formats[export_result.export_format]
        
        st.download_button(
            label=f"📥 Download {format_info['name']} ({export_result.file_size:,} bytes)",
            data=export_result.download_data,
            file_name=export_result.filename,
            mime=format_info['mime_type'],
            help=f"Download {export_result.items_exported} items as {format_info['name']}"
        )
    
    def get_export_preview(self, content_items: List[Any], export_format: str,
                          max_items: int = 3) -> str:
        """Generate a preview of how content will look in export format"""
        
        preview_items = content_items[:max_items]
        
        try:
            if export_format == "json":
                preview_data = self._export_json(preview_items, include_metadata=False)
                return preview_data.decode('utf-8')[:1000] + "..." if len(preview_data) > 1000 else preview_data.decode('utf-8')
            
            elif export_format == "markdown":
                preview_data = self._export_markdown(preview_items, include_metadata=False)
                return preview_data.decode('utf-8')[:1000] + "..." if len(preview_data) > 1000 else preview_data.decode('utf-8')
            
            elif export_format == "txt":
                preview_data = self._export_txt(preview_items)
                return preview_data.decode('utf-8')[:1000] + "..." if len(preview_data) > 1000 else preview_data.decode('utf-8')
            
            elif export_format == "csv":
                preview_data = self._export_csv(preview_items)
                return preview_data.decode('utf-8')[:1000] + "..." if len(preview_data) > 1000 else preview_data.decode('utf-8')
            
            else:
                return f"Preview not available for {export_format} format"
                
        except Exception as e:
            return f"Error generating preview: {str(e)}"
    
    def estimate_export_size(self, content_items: List[Any], export_format: str) -> Dict[str, Any]:
        """Estimate the size and characteristics of an export"""
        
        # Sample estimation based on first few items
        sample_size = min(len(content_items), 3)
        sample_items = content_items[:sample_size]
        
        try:
            if export_format == "json":
                sample_data = self._export_json(sample_items)
            elif export_format == "markdown":
                sample_data = self._export_markdown(sample_items)
            elif export_format == "txt":
                sample_data = self._export_txt(sample_items)
            elif export_format == "csv":
                sample_data = self._export_csv(sample_items)
            else:
                sample_data = b"Sample data"
            
            sample_size_bytes = len(sample_data)
            estimated_total_size = (sample_size_bytes * len(content_items)) // sample_size
            
            return {
                "estimated_size_bytes": estimated_total_size,
                "estimated_size_mb": estimated_total_size / (1024 * 1024),
                "sample_size_bytes": sample_size_bytes,
                "compression_recommended": estimated_total_size > 1024 * 1024,  # > 1MB
                "format_efficiency": "high" if export_format in ["json", "csv"] else "medium"
            }
            
        except Exception as e:
            return {
                "estimated_size_bytes": 0,
                "estimated_size_mb": 0,
                "error": str(e)
            }

